"""Generate Word (.docx) reports on demand from live, role-scoped data."""

from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.models import AiUsage, Lesson, Progress, School, SecurityLog, User
from app.models.enums import Role, SecurityStatus, UserStatus
from app.services.ai_usage import (
    usage_breakdown_for_school,
    usage_by_user,
    usage_total_for_school,
)

BRAND = RGBColor(0x0F, 0x76, 0x6E)  # teal-700
MUTED = RGBColor(0x64, 0x74, 0x8B)  # slate-500


def _heading(doc: Document, text: str, size: int = 14) -> None:
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = BRAND


def _meta_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not rows:
        _meta_line(doc, "No records.")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)


def _title_block(doc: Document, title: str, subtitle: str, generated_by: str) -> None:
    t = doc.add_paragraph()
    r = t.add_run("IM-Telligence")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BRAND

    s = doc.add_paragraph()
    sr = s.add_run(title)
    sr.bold = True
    sr.font.size = Pt(16)

    when = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    _meta_line(doc, f"{subtitle}")
    _meta_line(doc, f"Generated {when} · by {generated_by}")
    doc.add_paragraph()


def _finish(doc: Document) -> io.BytesIO:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# --------------------------------------------------------------------------- #
def _school_sections(db: Session, doc: Document, school: School) -> None:
    teachers = list(
        db.scalars(
            select(User).where(User.school_id == school.id, User.role == Role.teacher)
        )
    )
    tids = [t.id for t in teachers]
    name_by_id = {t.id: t.name for t in teachers}
    progress = (
        list(db.scalars(select(Progress).where(Progress.teacher_id.in_(tids)))) if tids else []
    )
    lesson_ids = {p.lesson_id for p in progress}
    lessons = (
        {l.id: l.title for l in db.scalars(select(Lesson).where(Lesson.id.in_(lesson_ids)))}
        if lesson_ids
        else {}
    )
    logs = list(
        db.scalars(
            select(SecurityLog)
            .where(SecurityLog.school_id == school.id)
            .order_by(SecurityLog.timestamp.desc())
            .limit(50)
        )
    )

    usage = usage_by_user(db, tids)
    ai = usage_breakdown_for_school(db, school.id)

    active = sum(1 for t in teachers if t.status == UserStatus.active)
    avg = round(sum(p.percent_complete for p in progress) / len(progress)) if progress else 0
    late = sum(1 for p in progress if p.watchdog.value == "late" or p.status.value == "late")
    alerts = sum(1 for l in logs if l.status != SecurityStatus.ok)

    _heading(doc, "Summary")
    _table(
        doc,
        ["Active teachers", "Assignments", "Avg completion", "Late", "Security alerts", "AI (teachers)", "AI (admin)"],
        [[str(active), str(len(progress)), f"{avg}%", str(late), str(alerts), str(ai["teacher"]), str(ai["admin"])]],
    )

    _heading(doc, "Teachers")
    _table(
        doc,
        ["Name", "Grades", "Language", "Status", "AI questions"],
        [
            [
                t.name,
                ", ".join(t.grades or []) or "—",
                (t.language or "—"),
                t.status.value,
                str(usage.get(t.id, {}).get("total", 0)),
            ]
            for t in teachers
        ],
    )

    _heading(doc, "AI assistant usage")
    _meta_line(doc, "Teacher questions to the lesson assistant — last 7 days and all time.")
    _meta_line(
        doc,
        f"School total {ai['total']}: {ai['teacher']} from teachers (below) + "
        f"{ai['admin']} from the school admin's operations assistant.",
    )
    _table(
        doc,
        ["Teacher", "Last 7 days", "Total questions"],
        [
            [
                t.name,
                str(usage.get(t.id, {}).get("last7", 0)),
                str(usage.get(t.id, {}).get("total", 0)),
            ]
            for t in sorted(
                teachers,
                key=lambda t: usage.get(t.id, {}).get("total", 0),
                reverse=True,
            )
        ],
    )

    _heading(doc, "Teacher progress")
    _table(
        doc,
        ["Teacher", "Lesson", "Status", "%", "Watchdog"],
        [
            [
                name_by_id.get(p.teacher_id, p.teacher_id),
                lessons.get(p.lesson_id, p.lesson_id),
                p.status.value,
                str(p.percent_complete),
                p.watchdog.value,
            ]
            for p in progress
        ],
    )

    _heading(doc, "Security log")
    _table(
        doc,
        ["User", "Event", "Status", "Device", "Time"],
        [
            [
                l.user_name,
                l.event.value,
                l.status.value,
                (l.device or "")[:40],
                l.timestamp.strftime("%Y-%m-%d %H:%M"),
            ]
            for l in logs
        ],
    )


def _clean_inline(text: str) -> str:
    # Strip the lightweight markdown emphasis the model tends to add.
    return text.replace("**", "").replace("__", "").strip()


def _render_narrative(doc: Document, narrative: str) -> None:
    """Render the assistant's markdown-ish narrative into Word paragraphs:
    `##`/`#` lines become headings, `-`/`*`/`•` lines become bullets, the rest
    are plain paragraphs."""
    for raw in narrative.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("#"):
            _heading(doc, _clean_inline(line.lstrip("#")), size=13)
        elif line[:2] in ("- ", "* ") or line.startswith("•"):
            doc.add_paragraph(_clean_inline(line.lstrip("-*• ")), style="List Bullet")
        else:
            p = doc.add_paragraph()
            run = p.add_run(_clean_inline(line))
            run.font.size = Pt(10)


def build_school_ai_report(
    db: Session, school_id: str, generated_by: str, narrative: str
) -> tuple[io.BytesIO, str]:
    """School report that leads with an AI-written executive narrative, followed
    by the same live data tables as the standard school report."""
    school = db.get(School, school_id)
    school_name = school.name if school else "School"
    doc = Document()
    _title_block(doc, "School Report", school_name, generated_by)

    _heading(doc, "Executive summary")
    _meta_line(doc, "AI-generated narrative based on your school's live data.")
    _render_narrative(doc, narrative)
    doc.add_paragraph()

    if school:
        _school_sections(db, doc, school)
    date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    return _finish(doc), f"IM-Telligence AI Report - {school_name} - {date}.docx"


def build_school_report(db: Session, school_id: str, generated_by: str) -> tuple[io.BytesIO, str]:
    school = db.get(School, school_id)
    school_name = school.name if school else "School"
    doc = Document()
    _title_block(doc, "School Report", school_name, generated_by)
    if school:
        _school_sections(db, doc, school)
    date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    return _finish(doc), f"IM-Telligence Report - {school_name} - {date}.docx"


def build_super_report(
    db: Session, generated_by: str, school_id: str | None = None
) -> tuple[io.BytesIO, str]:
    date = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    # Scoped to one school -> same as the school report.
    if school_id:
        return build_school_report(db, school_id, generated_by)

    # Platform-wide overview.
    doc = Document()
    _title_block(doc, "Platform Report", "All schools", generated_by)

    schools = list(db.scalars(select(School).order_by(School.name)))
    teachers = list(db.scalars(select(User).where(User.role == Role.teacher)))
    teacher_count = len(teachers)
    lesson_count = db.scalar(select(func.count(Lesson.id))) or 0

    all_progress = (
        list(db.scalars(select(Progress).where(Progress.teacher_id.in_([t.id for t in teachers]))))
        if teachers
        else []
    )
    assignments = len(all_progress)
    avg_completion = (
        round(sum(p.percent_complete for p in all_progress) / assignments) if assignments else 0
    )
    late_total = sum(
        1 for p in all_progress if p.watchdog.value == "late" or p.status.value == "late"
    )
    alerts_total = db.scalar(
        select(func.count(SecurityLog.id)).where(SecurityLog.status != SecurityStatus.ok)
    ) or 0
    ai_total = db.scalar(select(func.count(AiUsage.id))) or 0

    _heading(doc, "Platform summary")
    _table(
        doc,
        ["Schools", "Teachers", "Lessons", "Assignments"],
        [[str(len(schools)), str(teacher_count), str(lesson_count), str(assignments)]],
    )
    _table(
        doc,
        ["Avg completion", "Late lessons", "Security alerts", "AI interactions"],
        [[f"{avg_completion}%", str(late_total), str(alerts_total), str(ai_total)]],
    )

    # Per-school rollups for the overview table.
    progress_by_school: dict[str, list[Progress]] = {s.id: [] for s in schools}
    school_of_teacher = {t.id: t.school_id for t in teachers}
    for p in all_progress:
        sid = school_of_teacher.get(p.teacher_id)
        if sid in progress_by_school:
            progress_by_school[sid].append(p)

    _heading(doc, "Schools overview")
    rows = []
    for s in schools:
        sp = progress_by_school.get(s.id, [])
        s_teachers = sum(1 for t in teachers if t.school_id == s.id)
        s_avg = round(sum(p.percent_complete for p in sp) / len(sp)) if sp else 0
        s_ai = usage_total_for_school(db, s.id)
        rows.append(
            [s.name, s.city or "—", s.country or "—", str(s_teachers), f"{s_avg}%", str(s_ai)]
        )
    _table(doc, ["School", "City", "Country", "Teachers", "Avg %", "AI"], rows)

    # Per-school detail sections (each includes per-teacher AI usage).
    for s in schools:
        doc.add_page_break()
        _heading(doc, s.name, size=16)
        _meta_line(doc, f"{s.city or '—'}, {s.country or '—'}")
        _school_sections(db, doc, s)

    return _finish(doc), f"IM-Telligence Platform Report - {date}.docx"
